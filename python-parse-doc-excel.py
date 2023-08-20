
import logging
import tkinter as tk
from tkinter import filedialog
import os
from docx import Document
import re
import json
import docx
from openpyxl import Workbook
import pandas as pd

import glob

keywords = ['会议时间','会议地点', '会议名称', '参加会议人员']
map_array = {
    '会议时间': r"(\d{4}年\d{1,2}月\d{1,2}日\d{1,2}时\d{1,2}分)",
    '会议地点' : r"会议地点：(.*)",
    '会议名称' : r"会议名称：(.*)",
    '参加会议人员' : r"参加会议人员：(.*)"
}
def get_docx_files(directory):
    # 切换到指定目录
    os.chdir(directory)
    
    # 使用 glob 匹配所有 .docx 文件
    docx_files = glob.glob("*.docx")
    return docx_files

def open_directory():
    root = tk.Tk()
    root.withdraw()
    directory = filedialog.askdirectory()
    if directory:
        print("选择的文件夹路径", directory)
        #parse_docx_files(directory)
        return directory
    # 获取目录下的所有.docx文件
    #docx_files = glob.glob(os.path.join(directory, "*.docx"))
    

    
def extract_text_from_docx(docx_file):
    doc = docx.Document(docx_file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def extract_keywords_from_text(text, keywords):
    extracted_data = []
    for keyword in keywords:
        if keyword in text:
            extracted_data.append(text)
    return extracted_data


def extract_text_after_keyword(text, keyword):
    pattern = map_array[keyword]
   # print(pattern)
    extracted_data = {}
    match = re.search(pattern, text)
    if match:
        content = match.group(1)
        #print(keyword, time)
        extracted_data[keyword] = content
        return extracted_data
    else:
        return "None"

def get_dic_data():
    directory = open_directory()
    docx_files = get_docx_files(directory)
    all_data = {}
    if len(docx_files) > 0:
        for file in docx_files:
            print("doc文件：", file)
            text= extract_text_from_docx(file)

            extracted_data = extract_keywords_from_text(text, keywords)
            #print(extracted_data)

            for keyword in keywords:
                data = extract_text_after_keyword(text, keyword)
                if data != "None":
                    if keyword in all_data:
                        all_data[keyword].append(data[keyword])
                    else :
                        all_data[keyword] = [data[keyword]]
                #print(data)
                # wb = Workbook()
                # ws = wb.active

                # header = list(data.keys())
                # ws.append(header)
                # for row in zip(*data.values()):
                #     ws.append(row)
                
                # wb.save('output.xlsx')
    
            

    else:
        print("目录中没有 .docx 文件。")

    return all_data

def main():
    data = get_dic_data()
    wb = Workbook()
    ws = wb.active

    header = list(data.keys())
    ws.append(header)
    for row in zip(*data.values()):
        ws.append(row)
                    
    wb.save('output.xlsx')
                    
    #wb.save('output.xlsx')

if __name__ == "__main__":
    main()
# def extract_text_from_docx(docx_file):
#     doc = docx.Document(docx_file)
#     text = ""
#     for paragraph in doc.paragraphs:
#         text += paragraph.text + "\n"
#     return text

# def extract_keywords_from_text(text, keywords):
#     extracted_data = []
#     for keyword in keywords:
#         if keyword in text:
#             extracted_data.append(keyword)
#     return extracted_data

# def write_to_excel(extracted_data, excel_file):
#     wb = Workbook()
#     ws = wb.active
#     for i, data in enumerate(extracted_data, start=1):
#         ws.cell(row=i, column=1, value=data)
#     wb.save(excel_file)


# # 设置关键字和文件路径
# keywords = ["会议时间", "会议地点", "keyword3"]
# docx_file = "path/to/docx/file.docx"
# excel_file = "path/to/excel/file.xlsx"
# # 提取docx文件中的文字
# text = extract_text_from_docx(docx_file)

# # 根据关键字提取数据
# extracted_data = extract_keywords_from_text(text, keywords)
# # current_path = os.getcwd()
# # print(current_path)
# # json_file_path =os.path.join(current_path, 'config.json')

# # with open(json_file_path) as f:
# #     config_data = json.load(f)

# # config_dict = dict(config_data)

# # print(config_dict['会议时间'])

# # 定义一个空的字典
# # map_array = {
# #   #  '会议时间': r"(\d{4}年\d{1,2}月\d{1,2}日\d{1,2}时\d{1,2}分)",
# #     '会议地点': r"建德.*"
# # }

# # keywords = ["会议时间", "会议地点"]


# # # 打印整个字典
# # #print(map_array)

# # keywords = ['会议地点']

# # # 配置日志输出的格式和级别
# # logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')

# # def extract_text_after_keyword(text, keyword):
# #     pattern = map_array[keyword]
# #     print(pattern)
# #     match = re.search(pattern, text)
# #     if match:
# #         time = match.group(1)
# #         print(keyword, time)
# #         return time
# #     else:
# #         return "None"

# def open_directory():
#     root = tk.Tk()
#     root.withdraw()
#     directory = filedialog.askdirectory()
#     if directory:
#         print("选择的文件夹路径", directory)
#         parse_docx_files(directory)





# def parse_docx_files(directory):
#     for filename in os.listdir(directory):
#         if filename.endswith(".docx"):
#             file_path = os.path.join(directory, filename)
#             document = Document(file_path)
#             for paragraph in document.paragraphs:
#                 #logging.info(paragraph.text)
#                 text = str(paragraph.text)
#                 for keyword in keywords:
#                     result = extract_keywords_from_text(text, keyword)
#                     if result != "None" :
#                         print("result = ",result)

#             # 在这里进行你的关键字提取操作
#             # 例如，可以使用 document.paragraphs 获取段落内容
#             # 然后使用正则表达式或其他方法提取关键字

# # # 打开目录选择对话框
# # open_directory()

